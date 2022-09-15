from docx import Document
from pathlib import Path
from preprocessing import textract_processing as txt
from text_processing import tf_idf_cosine_similarity as tf_idf
import os
from pyresparser import ResumeParser
from docx import Document
from pathlib import Path
import numpy as np

cache_file = "C:\\Users\\vivek\\Desktop\\resume_ranking\\cached_files\\" 
document = Document()
req_document = "C:\\Users\\vivek\\Desktop\\resume_ranking\\demo_samples\\jd_primary_sec\\Java Job Description - IGreenData_080422.pdf"

def save_file(text, file, skill):
    document = Document()
    document.add_paragraph(text)
    
    document.save(cache_file + file + "_" + skill + ".docx")

fileName = Path(req_document).stem
job_desc = txt.get_content_as_string(req_document)

primSkillStart = job_desc.find("primary skills")
primSkillEnd = job_desc.find("secondary skills")
primary_skills = job_desc[primSkillStart + len("primary skills") + 1: primSkillEnd]
save_file(primary_skills, fileName, "primary")
primaryFile = cache_file + fileName + "_primary.docx"

prim_data = ResumeParser(primaryFile).get_extracted_data()
# sec_data = ResumeParser(secondaryFile).get_extracted_data()
if prim_data['skills'] is not None:
   primary_skills = " ".join(prim_data['skills'])

resume_doc_text = []
doct = "C:\\Users\\vivek\\Desktop\\resume_ranking\\demo_samples\\final_pres\\AiswaryaRamachandran.pdf"

data = ResumeParser(doct).get_extracted_data()
if data['skills'] is not None:
   resume_doc_text.append(" ".join(data['skills']))

prim = tf_idf.get_tf_idf_cosine_similarity(primary_skills,resume_doc_text)
cos_sim_prim = [2*ele for ele in prim]

# sec = tf_idf.get_tf_idf_cosine_similarity(secondary_skills,resume_doc_text)
sec = [0 for i in range(len(prim))]

norm = np.array([cos_sim_prim[i] + sec[i] for i in range(len(prim))])
cos_sim_arr = (norm - np.min(norm))/ (np.max(norm) - np.min(norm))
cos_sim_list = cos_sim_arr.tolist()


print(cos_sim_list)